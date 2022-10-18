# _*_coding:utf-8 _*_
# v1.2.2
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx import document
from docx.document import Document
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import datetime
from zhdate import ZhDate as lunar_date
import spider.getCity_Weather as getCity_Weather
import spider.Nanyang as Nanyang
import os,subprocess


class Weathers():
    def __init__(self):

        self.remind = str()  # 温馨提示
        self.clothes = str()  # 穿衣指数

        self.doc_name = '某年某月某日晚间气象预报'
        self.doc_path = r'C:\Users\Public\Desktop\晚间气象预报'

        self.wea = '未知'
        self.temLow = '未知'
        self.temHigh = '未知'

        # 获取天气
        # self.getweather()
        # 获取南阳市的紫外线强度和空气质量
        # self.rays = Nanyang.get_rays()
        # self.airs = Nanyang.get_airs()
        # self.get_docx()

    def set_text(self):  # 冗余控制，避免出现误操作
        # 意义不大
        # 判断输入内容
        if not len(self.clothes):
            self.clothes = '________'
        if not len(self.remind):
            self.remind = '________________________________________________________________________'
        if not len(self.doc_path):
            self.doc_path = r'C:\Users\Public\Desktop\晚间气象预报'

    def getweather(self):
        self.airs = Nanyang.get_airs()
        self.rays = Nanyang.get_rays()
        href = "http://www.weather.com.cn/weather/101180701.shtml"
        weatherData = getCity_Weather.Weather(url=href)
        weatherData.url = href
        weatherData.getWeather()
        self.wea = weatherData.wea
        self.temHigh = weatherData.temHigh
        self.temLow = weatherData.temLow
        self.weatherList = weatherData.citys_weather()

    def getData(self):
        # 设置日期信息
        self.year = str(datetime.datetime.now().year)
        self.month = str(datetime.datetime.now().month)
        self.day = str(datetime.datetime.now().day)
        date = str(lunar_date.today().chinese())  # 今天的农历日期（中文格式）
        # 将日期格式剪切修改
        self.date = date[5:-9]
        # 将星期1转化成星期一
        weekday_list = ['一', '二', '三', '四', '五', '六', '日']
        self.weekday = str(weekday_list[datetime.datetime.today().weekday()])

    def get_docx(self):  # 生成文本信息
        self.set_text()  # 配置输入内容

        head = '晚间气象预报'
        paragraph_1 = f'亲爱的同学们，大家晚上好，欢迎收听晚间气象预报，今天是\n公历{self.month} 月 {self.day} 日\n农历 {self.date}\t\t\t\t\t\t\t\t星期{self.weekday}'
        paragraph_2 = '\n\n'
        paragraph_3 = f'接下来让我们关注一下南阳市明天白天的天气情况，预计南阳市明天白天{self.wea}，最低温度{self.temLow}，最高温度{self.temHigh}'
        paragraph_4 = f'\n紫外线强度 {self.rays}\t\t\t\t\t空气质量 {self.airs}\n穿衣指数{self.clothes}% \n在此温馨提醒同学们:\n{self.remind}\n'
        paragraph_5 = f'下面让我们关注一下其他城市的天气情况：\n\n{self.weatherList}\n\n好了，亲爱的同学们，晚间气象预报为您播送完了，感谢您的收听，再见。'

        text = paragraph_1+paragraph_2+paragraph_3+paragraph_4+paragraph_5  # 拼接所有文本
        # 创建docx文档
        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        p_head = doc.add_paragraph()
        p_head.space_before = Pt(40)
        p_head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 添加文字块
        head_text = p_head.add_run(head)
        head_text.bold = True
        head_text.font.color.rgb = RGBColor(0, 0, 0)
        head_text.font.size = Pt(16)

        p_weather = doc.add_paragraph()
        p_weather.space_before = Pt(28)

        run = p_weather.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(14)

        # 设置文件名
        self.doc_name = self.year+self.month+self.day+'晚间气象预报.docx'

        if not os.path.exists(self.doc_path):
            os.makedirs(self.doc_path)

        doc.save(f'{self.doc_path}/{self.doc_name}')
        os.startfile(f'{self.doc_path}/{self.doc_name}')

