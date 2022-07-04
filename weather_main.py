# _*_coding:utf-8 _*_
# v1.2.2
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx import document
import tkinter as tk
from docx.document import Document
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import datetime
from zhdate import ZhDate as lunar_date
import getCity_Weather
import Nanyang
import tkinter.messagebox
import os


class Application(tk.Frame):
    def __init__(self, master=None):
        super().__init__(master)
        self.master = master
        self.pack()

        self.remind = '________________________________________________________________________'  # 温馨提示
        self.clothes = '________'  # 穿衣指数

        self.doc_name = '某年某月某日晚间气象预报'
        self.doc_path = 'D:/晚间气象预报/'

        self.wea = '未知'
        self.temLow = '未知'
        self.temHigh = '未知'

        # 获取天气
        self.getweather()

        # 获取南阳市的紫外线强度和空气质量
        self.rays = Nanyang.get_rays()
        self.airs = Nanyang.get_airs()

        # 创建可视化窗口
        self.createWidget()

    def set_text(self):  # 冗余控制，避免出现误操作
        # 意义不大
        # 判断输入内容
        if len(self.clothes_text.get()) == 0 or self.clothes_text.get() == '默认为下划线':
            self.clothes = '________'
        else:
            self.clothes = self.clothes_text.get()

        if len(self.remind_text.get()) == 0 or self.remind_text.get() == '默认为下划线':
            self.remind = '________________________________________________________________________'
        else:
            self.remind = self.remind_text.get()

        if self.doc_path_text.get() == '默认保存地址' or self.doc_path_text.get() == '':
            self.doc_path = 'D:/晚间气象预报/'
        else:
            self.doc_path = self.doc_path_text.get()

    def getweather(self):
        href = "http://www.weather.com.cn/weather/101180701.shtml"
        weatherData = getCity_Weather.Weather(url=href)
        weatherData.url = href
        weatherData.getWeather()
        self.wea = weatherData.wea
        self.temHigh = weatherData.temHigh+'℃'
        self.temLow = weatherData.temLow
        self.weatherList = weatherData.citys_weather()

    def doc_Scr(self):
        # 返回生成文件的地址
        return self.doc_path+self.doc_name

    def createWidget(self):  # 控制布局
        self.label_01 = tk.Label(self, text='信息填写完毕后，点击“生成文件”按钮即可生成今晚所需的天气预报电子稿件', width=16, font=(
            "宋体", 16), background='red', anchor='center')
        self.label_01.grid(ipadx=280, padx=4, pady=4,
                           row=0, column=0, columnspan=3)

        # 输入穿衣指数
        label_03 = tk.Label(self, text='输入穿衣指数', width=16,
                            font=("宋体", 16), relief='flat')
        label_03.grid(ipadx=60, padx=4, pady=4, row=2, column=0)

        self.clothes_text = tk.StringVar()
        self.clothes_text.set('默认为下划线')
        clos = tk.Entry(self, textvariable=self.clothes_text, font=("宋体", 14))
        clos.grid(padx=4, pady=4, row=2, column=1)

        # 输入温馨提示
        label_04 = tk.Label(self, text='输入温馨提醒', width=16,
                            font=("宋体", 16), relief='flat')
        label_04.grid(ipadx=60, padx=4, pady=4, row=3, column=0)

        self.remind_text = tk.StringVar()
        self.remind_text.set('默认为下划线')
        rmid = tk.Entry(self, textvariable=self.remind_text, font=("宋体", 14))
        rmid.grid(padx=4, pady=4, row=3, column=1)

        # 输入文件归档地址
        self.doc_path_text = tk.StringVar()
        self.doc_path_text.set('默认保存地址')
        document_scr = tk.Entry(
            self, textvariable=self.doc_path_text, font=("宋体", 14))
        document_scr.grid(padx=4, pady=4, row=4, column=0)

        # 生成文件按钮
        self.btn_01 = tk.Button(self, text='生成文件', width=16, font=(
            "宋体", 16), relief='flat', background='#c0c0c0', command=self.getfile)
        self.btn_01.grid(padx=4, pady=4, row=4, column=1)

        # 显示南阳市天气信息
        text = '紫外线强度：'+self.rays
        self.label_new_rays = tk.Label(
            self, text=text, width=16, font=("宋体", 16), anchor='center')
        self.label_new_rays.grid(padx=4, pady=4, row=5, column=0)

        text = '天气：'+self.wea
        self.label_new_wea = tk.Label(
            self, text=text, width=16, font=("宋体", 16), anchor='center')
        self.label_new_wea.grid(padx=4, pady=4, row=5, column=1)

        text = '最高气温：'+self.temHigh
        self.label_new_temHigh = tk.Label(
            self, text=text, width=16, font=("宋体", 16), anchor='center')
        self.label_new_temHigh.grid(padx=4, pady=4, row=6, column=0)

        text = '最低气温：'+self.temLow
        self.label_new_temLow = tk.Label(
            self, text=text, width=16, font=("宋体", 16), anchor='center')
        self.label_new_temLow.grid(padx=4, pady=4, row=6, column=1)

        # 作者信息
        self.label_02 = tk.Label(self, text='by Muyuan', width=16, font=(
            "宋体", 9), justify='right', anchor='center')
        self.label_02.grid(padx=4, pady=4, row=6, column=2)

        # 退出按钮
        self.btnQuit = tk.Button(self, text='退出', width=5, font=(
            "宋体", 16), relief='flat', background='red', command=self.master.destroy)
        self.btnQuit.grid(padx=4, pady=4, row=7, column=0, columnspan=2)

    def getfile(self):
        msg = tkinter.messagebox.askokcancel(
            title='提示', message='是否开始生成气象预报')  # 弹出窗口以便确认执行状态
        if msg == False:
            return False
        self.get_docx()
        return True

    def get_docx(self):  # 生成文本信息
        self.set_text()  # 配置输入内容
        # 设置日期信息
        year = str(datetime.datetime.now().year)
        month = str(datetime.datetime.now().month)
        day = str(datetime.datetime.now().day)
        date = str(lunar_date.today().chinese())  # 今天的农历日期（中文格式）

        # 将日期格式剪切修改
        date = date[5:-9]

        # 将星期1转化成星期一
        weekday_list = ['一', '二', '三', '四', '五', '六', '日']
        weekday = str(weekday_list[datetime.datetime.today().weekday()])

        head = '晚间气象预报'
        paragraph_1 = '亲爱的同学们，大家晚上好，欢迎收听晚间气象预报，今天是\n公历'+month+' 月 '+day+' 日\n农历 '+date + \
            '\t\t\t\t\t\t\t\t星期'+weekday
        paragraph_2 = '\n\n'
        paragraph_3 = '接下来让我们关注一下南阳市明天白天的天气情况，预计南阳市明天白天  ' + \
            self.wea+' ，最低温度 '+self.temLow+'  最高温度 '+self.temHigh
        paragraph_4 = '\n紫外线强度 '+self.rays+'\t\t\t\t\t空气质量 '+self.airs+'\n穿衣指数' + \
            self.clothes+'% \n在此温馨提醒同学们:\n'+self.remind+'\n'
        paragraph_5 = '下面让我们关注一下其他城市的天气情况：\n\n'+self.weatherList + \
            '\n\n好了，亲爱的同学们，晚间气象预报为您播送完了，感谢您的收听，再见。'

        text = paragraph_1+paragraph_2+paragraph_3+paragraph_4+paragraph_5  # 拼接所有文本
        # 创建docx文档
        doc = Document()
        doc.styles['Normal'].font.name = u'宋体'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

        p_head = doc.add_paragraph()
        p_head.space_before = Pt(40)
        p_head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 添加文字块
        head_text = p_head.add_run(head)
        head_text.bold = True
        head_text.font.color.rgb = RGBColor(0, 0, 0)
        head_text.font.size = Pt(16)

        p_weather = doc.add_paragraph()
        p_weather.space_before = Pt(28)

        run = p_weather.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(14)

        # 设置文件名
        self.doc_name = year+month+day+'晚间气象预报.docx'

        if not os.path.exists(self.doc_path):
            os.makedirs(self.doc_path)

        doc.save(self.doc_Scr())

        mesg = self.doc_name+' 已在 '+self.doc_path+' 下生成'
        tkinter.messagebox.showinfo(
            title='成功', message=mesg)
        os.system(self.doc_Scr())


def __main__():
    root = tk.Tk()
    root.geometry("1000x300+500+300")
    root.title("气象预报生成器v1.2.1")
    application = Application(master=root)
    application.mainloop()


if __name__ == '__main__':
    __main__()
